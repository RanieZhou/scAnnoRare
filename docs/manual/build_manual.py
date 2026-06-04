# -*- coding: utf-8 -*-
"""
生成《单细胞细胞类型注释与稀有细胞识别多方法评估系统 V1.0 用户操作手册》
输出: docs/manual/scAnnoRare_用户操作手册_V1.0.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "screenshots")  # screenshots folder

SOFT_NAME = "单细胞细胞类型注释与稀有细胞识别多方法评估系统"
SOFT_VER = "V1.0"
HEADER_TEXT = f"{SOFT_NAME} {SOFT_VER}"

# ── colors ────────────────────────────────────────────────────────────────────
C_TITLE = RGBColor(0x33, 0x33, 0x33)
C_H1 = RGBColor(0x1F, 0x3A, 0x5F)
C_H2 = RGBColor(0x2C, 0x3E, 0x50)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_MUTE = RGBColor(0x66, 0x66, 0x66)

CN_FONT = "宋体"
CN_HEAD = "黑体"
EN_FONT = "Times New Roman"


# ── low-level helpers ─────────────────────────────────────────────────────────
def _set_run_font(run, cn=CN_FONT, en=EN_FONT, size=12, bold=False, color=C_BODY):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)
    rfonts.set(qn('w:eastAsia'), cn)


def _para(doc, text="", size=12, bold=False, cn=CN_FONT, en=EN_FONT,
          color=C_BODY, align=None, before=0, after=6, indent_first=True,
          line=20):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    if align is not None:
        p.alignment = align
    if indent_first and text:
        pf.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        _set_run_font(r, cn=cn, en=en, size=size, bold=bold, color=color)
    return p


def h1(doc, num, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(f"第 {num} 章  {text}")
    _set_run_font(r, cn=CN_HEAD, en=EN_FONT, size=18, bold=True, color=C_H1)
    # bookmark-like bottom border
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    _set_run_font(r, cn=CN_HEAD, en=EN_FONT, size=14, bold=True, color=C_H2)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_run_font(r, cn=CN_HEAD, en=EN_FONT, size=12.5, bold=True, color=C_H2)
    return p


def body(doc, text):
    return _para(doc, text, size=12, after=6, line=22)


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(24 + level * 18)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    r0 = p.add_run("● " if level == 0 else "○ ")
    _set_run_font(r0, size=10, color=C_MUTE)
    r = p.add_run(text)
    _set_run_font(r, size=12, color=C_BODY)
    return p


def step(doc, n, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(24)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(22)
    r0 = p.add_run(f"步骤 {n}　")
    _set_run_font(r0, cn=CN_HEAD, size=12, bold=True, color=C_H1)
    r = p.add_run(text)
    _set_run_font(r, size=12, color=C_BODY)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(18)
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    r0 = p.add_run("【提示】")
    _set_run_font(r0, cn=CN_HEAD, size=11, bold=True, color=RGBColor(0xB0, 0x60, 0x00))
    r = p.add_run(text)
    _set_run_font(r, size=11, color=C_MUTE)
    return p


def figure(doc, filename, caption, width_cm=14.5):
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        _set_run_font(r, cn=CN_FONT, size=10.5, color=C_MUTE)
    else:
        note(doc, f"（截图缺失：{filename}）")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        pp = hdr[i].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(htxt)
        _set_run_font(r, cn=CN_HEAD, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr[i], "2C3E50")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            pp = cells[i].paragraphs[0]
            pp.paragraph_format.space_after = Pt(2)
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pp.paragraph_format.line_spacing = Pt(18)
            r = pp.add_run(str(val))
            _set_run_font(r, size=10, color=C_BODY)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return t


def _shade_cell(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcpr.append(shd)


def code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Pt(24)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(18)
        r = p.add_run(ln if ln else " ")
        _set_run_font(r, cn=CN_FONT, en="Consolas", size=10,
                      color=RGBColor(0x10, 0x40, 0x10))
        _shade_para(p, "F2F4F5")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def _shade_para(p, hexcolor):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    ppr.append(shd)


# ── page number field ─────────────────────────────────────────────────────────
def _add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = paragraph.add_run("第 ")
    _set_run_font(r1, size=9, color=C_MUTE)

    def field(run_para, instr):
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr_el = OxmlElement('w:instrText'); instr_el.set(qn('xml:space'), 'preserve'); instr_el.text = instr
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run = run_para.add_run()
        run._r.append(fld_begin); run._r.append(instr_el); run._r.append(fld_end)
        _set_run_font(run, size=9, color=C_MUTE)

    field(paragraph, "PAGE")
    r2 = paragraph.add_run(" 页 / 共 ")
    _set_run_font(r2, size=9, color=C_MUTE)
    field(paragraph, "NUMPAGES")
    r3 = paragraph.add_run(" 页")
    _set_run_font(r3, size=9, color=C_MUTE)


def setup_section(section, with_header=True):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.4)
    section.footer_distance = Cm(1.2)
    if with_header:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(HEADER_TEXT)
        _set_run_font(hr, cn=CN_FONT, size=9, color=C_MUTE)
        _add_bottom_border(hp)
    fp = section.footer.paragraphs[0]
    _add_page_number(fp)


def _add_bottom_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
    pbdr.append(bottom)
    ppr.append(pbdr)


# import content builder
from manual_content import build_content


def main():
    doc = Document()

    # default style
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)

    helpers = dict(
        para=_para, h1=h1, h2=h2, h3=h3, body=body, bullet=bullet,
        step=step, note=note, figure=figure, table=table, code_block=code_block,
        set_run_font=_set_run_font, WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
        Pt=Pt, Cm=Cm, RGBColor=RGBColor, SOFT_NAME=SOFT_NAME, SOFT_VER=SOFT_VER,
        C_TITLE=C_TITLE, C_H1=C_H1, C_MUTE=C_MUTE, C_BODY=C_BODY,
        CN_HEAD=CN_HEAD, CN_FONT=CN_FONT,
    )

    build_content(doc, helpers)

    setup_section(doc.sections[0])

    out = os.path.join(HERE, "scAnnoRare_用户操作手册_V1.0.docx")
    doc.save(out)
    print("Saved:", out)
    print("Paragraphs:", len(doc.paragraphs))


if __name__ == "__main__":
    main()
